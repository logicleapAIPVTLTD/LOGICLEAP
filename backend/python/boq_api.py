# import re
# import os
# import sys
# import json
# import pandas as pd
# from docx import Document
# import pdfplumber
# from pdf2image import convert_from_path
# import pytesseract
# import cv2
# import numpy as np
# from PIL import Image
# import boto3
# from decimal import Decimal

# # AWS Configuration
# AWS_REGION = "ap-south-2"
# AWS_ACCESS_KEY_ID = "AKIAYH3VJY2ZUOPIZ27O"
# AWS_SECRET_ACCESS_KEY = "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"

# # Base directory setup
# BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# DATA_DIR = os.path.join(BASE_DIR, 'data')

# # Tesseract configuration
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# OCR_CONFIG = "--oem 3 --psm 6 -c preserve_interword_spaces=1"

# def dynamodb_table_to_df(table_name):
#     dynamodb = boto3.resource(
#         "dynamodb",
#         region_name=AWS_REGION,
#         aws_access_key_id=AWS_ACCESS_KEY_ID,
#         aws_secret_access_key=AWS_SECRET_ACCESS_KEY
#     )

#     table = dynamodb.Table(table_name)
#     items = []
#     scan_kwargs = {}

#     while True:
#         response = table.scan(**scan_kwargs)
#         items.extend(response.get("Items", []))

#         if "LastEvaluatedKey" in response:
#             scan_kwargs["ExclusiveStartKey"] = response["LastEvaluatedKey"]
#         else:
#             break

#     def convert(val):
#         if isinstance(val, Decimal):
#             return int(val) if val % 1 == 0 else float(val)
#         return val

#     for item in items:
#         for k, v in item.items():
#             item[k] = convert(v)

#     return pd.DataFrame(items)

# def ocr_image(image: Image.Image) -> str:
#     img = np.array(image)
#     gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
#     gray = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)[1]
#     return pytesseract.image_to_string(gray, config=OCR_CONFIG)

# def ocr_image_from_path(image_path: str) -> str:
#     if not os.path.exists(image_path):
#         raise FileNotFoundError("Image not found")

#     image = Image.open(image_path)
#     img = np.array(image)
#     gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
#     gray = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)[1]
#     return pytesseract.image_to_string(gray, config=OCR_CONFIG)

# def ocr_pdf(path: str) -> str:
#     # Update poppler path if needed
#     poppler_path = r"C:\Users\VICTUS\Downloads\Release-23.11.0-0\poppler-23.11.0\Library\bin"
    
#     pages = convert_from_path(path, dpi=300, poppler_path=poppler_path)
#     texts = []
#     for page in pages:
#         texts.append(ocr_image(page))
#     return "\n".join(texts)

# def extract_text(path: str) -> str:
#     ext = os.path.splitext(path)[1].lower()

#     if ext == ".txt":
#         return open(path, encoding="utf-8").read()

#     if ext == ".docx":
#         doc = Document(path)
#         text = "\n".join(p.text for p in doc.paragraphs)
#         if len(text.strip()) > 50:
#             return text
#         raise ValueError("DOCX seems scanned. Convert to PDF.")

#     if ext == ".pdf":
#         with pdfplumber.open(path) as pdf:
#             text = "\n".join(p.extract_text() or "" for p in pdf.pages)
#         if len(text.strip()) > 100:
#             return text
#         return ocr_pdf(path)

#     if ext in {".png", ".jpg", ".jpeg"}:
#         return ocr_image_from_path(path)

#     raise ValueError("Unsupported file format")

# def normalize_text(text: str) -> str:
#     text = text.lower()
#     corrections = {
#         "lenghth": "length",
#         "widht": "width",
#         "widhth": "width",
#         "hight": "height",
#         "wid": "width",
#         "len": "length"
#     }
#     for wrong, correct in corrections.items():
#         text = text.replace(wrong, correct)
#     text = re.sub(r"[^a-z0-9.\s]", " ", text)
#     text = re.sub(r"\s+", " ", text).strip()
#     return text

# def normalize(text: str) -> str:
#     text = text.lower()
#     text = re.sub(r"[^a-z0-9\s]", " ", text)
#     text = re.sub(r"\s+", " ", text).strip()
#     return text

# # Load master data
# work_df = dynamodb_table_to_df("work_master").fillna("")
# room_df = dynamodb_table_to_df("room_master").fillna("")

# work_df["kw_norm"] = work_df["keyword"].apply(normalize)
# work_df["family_norm"] = work_df["normalized_keyword"].apply(normalize)
# room_df["kw_norm"] = room_df["keyword"].apply(normalize)

# def split_sentences(text: str):
#     text = text.replace("\r", "")
#     parts = re.split(r"\n\s*\n|(?<!\d)\.(?!\d)", text)
#     return [p.strip() for p in parts if len(p.strip()) > 5]

# ACTION_VERBS = {
#     "provide", "providing", "fix", "fixing", "install", "installation", 
#     "installing", "lay", "laying", "erect", "erection", "construct", 
#     "construction", "execute", "execution", "carry", "carrying", 
#     "supply", "supplying", "apply", "applying", "replace", "replacement",
#     "remove", "removal", "clean", "cleaning",
# }

# def has_action_verb(sentence: str) -> bool:
#     tokens = normalize(sentence).split()
#     if not any(v in tokens for v in ACTION_VERBS):
#         return False
#     if re.search(r"\d", sentence) and not any(v in tokens for v in ACTION_VERBS):
#         return False
#     return True

# def is_work_sentence(sentence: str) -> bool:
#     text = normalize(sentence)
#     return any(v in text.split() for v in ACTION_VERBS)

# def merge_context(sentences):
#     merged = []
#     buffer = None

#     for s in sentences:
#         from boq_api import detect_all_works
#         works = detect_all_works(s)

#         if works:
#             if buffer:
#                 merged.append(buffer)
#             buffer = s
#             continue

#         if buffer:
#             buffer = buffer + " " + s

#     if buffer:
#         merged.append(buffer)

#     return merged

# def detect_room(text):
#     clean = normalize(text)
#     matches = []
    
#     for _, r in room_df.iterrows():
#         kw = r["kw_norm"]
#         if kw and kw in clean:
#             matches.append({
#                 "length": len(kw),
#                 "room_code": r["room_code"],
#                 "room_name": r["room_name"]
#             })

#     if not matches:
#         return "UNSPEC", "Unspecified"

#     matches.sort(key=lambda x: x["length"], reverse=True)
#     best = matches[0]
#     return best["room_code"], best["room_name"]

# def detect_all_works(sentence):
#     clean = normalize(sentence)
#     candidates = []
#     tokens = set(clean.split())

#     if "room" in tokens and any(t in tokens for t in {"clean", "cleaning", "cleaned"}):
#         return [{
#             "work_code": "SERV_ROOM_CLEAN",
#             "work_name": "Room Cleaning",
#             "score": 2000
#         }]

#     for _, w in work_df.iterrows():
#         kw = w["kw_norm"]
#         if not kw:
#             continue
#         if kw in clean:
#             candidates.append({
#                 "work_code": w["work_code"],
#                 "work_name": w["work_name"],
#                 "score": 1000 + len(kw)
#             })

#     if candidates:
#         best = max(candidates, key=lambda x: x["score"])
#         return [best]

#     phrase_hits = []
#     sorted_rows = sorted(
#         work_df.to_dict("records"),
#         key=lambda x: len(x["kw_norm"]),
#         reverse=True
#     )

#     for w in sorted_rows:
#         kw = w["kw_norm"]
#         if not kw:
#             continue
#         if kw in clean:
#             phrase_hits.append({
#                 "work_code": w["work_code"],
#                 "work_name": w["work_name"],
#                 "score": 800 + len(kw)
#             })

#     if phrase_hits:
#         best = max(phrase_hits, key=lambda x: x["score"])
#         return [best]

#     sent_tokens = set(clean.split())
#     for _, w in work_df.iterrows():
#         kw = w["kw_norm"]
#         if not kw:
#             continue
#         kw_tokens = set(kw.split())
#         hit = len(kw_tokens & sent_tokens)
#         if hit == 0:
#             continue
#         score = hit * 100 + len(kw_tokens) * 50 + (30 if kw in clean else 0)
#         candidates.append({
#             "work_code": w["work_code"],
#             "work_name": w["work_name"],
#             "score": score
#         })

#     if candidates:
#         best = max(candidates, key=lambda x: x["score"])
#         return [best]

#     return []

# def extract_dimensions(text: str):
#     text = normalize_text(text)
#     dims = {}

#     def add(key, value, unit):
#         dims.setdefault(key, []).append({
#             "value": round(value, 2) if isinstance(value, (int, float)) else value,
#             "unit": unit
#         })

#     length_match = re.search(r"(\d+(?:\.\d+)?)\s*(?:length|len)", text)
#     width_match = re.search(r"(\d+(?:\.\d+)?)\s*(?:width|wid)", text)

#     if length_match:
#         length_val = float(length_match.group(1))
#         add("length", length_val, "m")

#     if width_match:
#         width_val = float(width_match.group(1))
#         add("width", width_val, "m")

#     if length_match and width_match:
#         area_val = length_val * width_val
#         add("area", area_val, "sqm")

#     for m in re.finditer(r"(\d+(?:\.\d+)?)\s*(sqft|sq\.?ft|sft|sqm|sq\.?m|m2)", text):
#         add("area", float(m.group(1)), m.group(2))

#     for m in re.finditer(r"(\d+(?:\.\d+)?)\s*x\s*(\d+(?:\.\d+)?)\s*(mm|cm|m)\b", text):
#         add("dimension_pair", f"{m.group(1)}x{m.group(2)}", m.group(3))

#     for m in re.finditer(r"(\d+(?:\.\d+)?)\s*\b(m|meter|metre|ft|feet|rm|rmt)\b", text):
#         add("length", float(m.group(1)), m.group(2))

#     for m in re.finditer(r"(height|width|depth|thickness)\s*(of)?\s*(\d+(?:\.\d+)?)\s*(mm|cm|m|ft)", text):
#         add(m.group(1), float(m.group(3)), m.group(4))

#     for m in re.finditer(r"(\d+(?:\.\d+)?)\s*(cum|cu\.?m|m3)", text):
#         add("volume", float(m.group(1)), m.group(2))

#     for m in re.finditer(r"(\d+)\s*(nos|no|numbers|pcs|pieces|units|sets|points)", text):
#         add("quantity", int(m.group(1)), m.group(2))

#     return dims

# def compute_confidence(room_name, dims, work_name):
#     score = 0.6
#     if room_name != "Unspecified":
#         score += 0.15
#     if dims:
#         score += 0.15
#     if len(work_name.split()) >= 2:
#         score += 0.1
#     return round(min(score, 0.95), 2)

# def text_to_boq(text, min_confidence=0.45):
#     sentences = split_sentences(text)
#     blocks = merge_context(sentences)
#     results = []
#     seen = set()

#     for block in blocks:
#         room_code, room_name = detect_room(block)
#         works = detect_all_works(block)
#         dims = extract_dimensions(block)

#         for w in works:
#             if room_code == "UNSPEC":
#                 dedup_key = w["work_code"]
#             else:
#                 dedup_key = (room_code, w["work_code"])

#             if dedup_key in seen:
#                 continue
#             seen.add(dedup_key)

#             confidence = compute_confidence(room_name, dims, w["work_name"])

#             if confidence < min_confidence:
#                 continue

#             results.append({
#                 "room_name": room_name,
#                 "work_code": w["work_code"],
#                 "work_name": w["work_name"],
#                 "dimensions": dims,
#                 "confidence": confidence,
#                 "source": "TEXT_INPUT",
#                 "context": block
#             })

#     return results

# def main():
#     """CLI entry point for API calls"""
#     try:
#         if len(sys.argv) < 2:
#             print(json.dumps({"error": "No input provided"}))
#             sys.exit(1)

#         input_type = sys.argv[1]

#         if input_type == "text":
#             text = sys.argv[2]
#         elif input_type == "file":
#             file_path = sys.argv[2]
#             if not os.path.exists(file_path):
#                 print(json.dumps({"error": f"File not found: {file_path}"}))
#                 sys.exit(1)
#             text = extract_text(file_path)
#         else:
#             print(json.dumps({"error": "Invalid input type"}))
#             sys.exit(1)

#         min_confidence = float(sys.argv[3]) if len(sys.argv) > 3 else 0.45
#         boq = text_to_boq(text, min_confidence)
        
#         print(json.dumps({
#             "success": True,
#             "data": boq,
#             "count": len(boq)
#         }))

#     except Exception as e:
#         print(json.dumps({
#             "success": False,
#             "error": str(e)
#         }))
#         sys.exit(1)

# if __name__ == "__main__":
#     main()


#!/usr/bin/env python3
"""
BOQ API Integration Script
Handles stdin/stdout communication with Node.js controller
Supports: Text, Document, and Image processing modes
"""

import sys
import json
import os
import re
import boto3
import pandas as pd
import pdfplumber
from docx import Document
from decimal import Decimal
import google.generativeai as genai
from difflib import get_close_matches

# =========================================================
# 1. CONFIGURATION
# =========================================================

# AWS Credentials
AWS_REGION = "ap-south-2"
AWS_ACCESS_KEY_ID = "AKIAYH3VJY2ZUOPIZ27O"
AWS_SECRET_ACCESS_KEY = "bj/52jjCrCg3yYAcPOMZdCVG+OYqHeIin+fXFiKm"

# Gemini Config
GEMINI_API_KEY = "AIzaSyAi3-skaFvtlgA8UTofE6JkQgCKmcAKZ_4"  # Replace with your key
GEMINI_MODEL_NAME = "gemini-2.5-flash"

# DynamoDB Tables
TABLE_WORK_MASTER = "work_master"
TABLE_ROOM_MASTER = "room_master"

genai.configure(api_key=GEMINI_API_KEY)

# =========================================================
# 2. DYNAMODB LIVE CONNECTION
# =========================================================

def get_dynamodb_data(table_name):
    """Fetch all data from DynamoDB table"""
    try:
        dynamodb = boto3.resource(
            "dynamodb",
            region_name=AWS_REGION,
            aws_access_key_id=AWS_ACCESS_KEY_ID,
            aws_secret_access_key=AWS_SECRET_ACCESS_KEY
        )
        table = dynamodb.Table(table_name)
        
        items = []
        response = table.scan()
        items.extend(response.get("Items", []))
        
        while "LastEvaluatedKey" in response:
            response = table.scan(ExclusiveStartKey=response["LastEvaluatedKey"])
            items.extend(response.get("Items", []))

        def convert(val):
            if isinstance(val, Decimal):
                return int(val) if val % 1 == 0 else float(val)
            return val

        clean_items = [{k: convert(v) for k, v in i.items()} for i in items]
        return pd.DataFrame(clean_items)

    except Exception as e:
        sys.stderr.write(f"❌ Database Connection Failed: {e}\n")
        return pd.DataFrame()

# LOAD DATA GLOBALLY
work_df = get_dynamodb_data(TABLE_WORK_MASTER).fillna("")
room_df = get_dynamodb_data(TABLE_ROOM_MASTER).fillna("")

if not work_df.empty:
    WORK_NAMES = work_df['work_name'].tolist()
    WORK_MAP = dict(zip(work_df['work_name'], work_df['work_code']))
else:
    WORK_NAMES = []
    WORK_MAP = {}

# =========================================================
# 3. TEXT PROCESSING ENGINE
# =========================================================

def normalize_text(text: str):
    """Normalize text for consistent processing"""
    text = text.lower()
    replacements = {
        "sq.ft": "sqft", "square feet": "sqft", "sq m": "sqm",
        "square meters": "sqm", "square metres": "sqm",
        "rft": "rft", "running feet": "rft", "running meter": "rmt"
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"[^a-z0-9.x\s]", " ", text)
    return re.sub(r"\s+", " ", text).strip()

def extract_metrics(text: str):
    """Extract all dimensional information from text"""
    text = normalize_text(text)
    dims = {}

    def add_dim(key, val, unit):
        try:
            final_val = float(val)
        except ValueError:
            final_val = val
        dims.setdefault(key, []).append({"value": final_val, "unit": unit})

    # Area patterns
    for m in re.finditer(r"(\d+(?:\.\d+)?)\s*(sqft|sqm|m2)", text):
        add_dim("area", m.group(1), m.group(2))
    
    # Dimensions (e.g., 600x600)
    for m in re.finditer(r"(\d+(?:\.\d+)?)\s*x\s*(\d+(?:\.\d+)?)\s*(m|ft|mm)?", text):
        full_dim = f"{m.group(1)}x{m.group(2)}"
        add_dim("dims", full_dim, m.group(3) or "mm")

    # Running length
    for m in re.finditer(r"(\d+(?:\.\d+)?)\s*(rft|rmt|meter|m)", text):
        add_dim("length", m.group(1), m.group(2))

    # Height/Width/Depth
    for m in re.finditer(r"(height|width|depth|thickness)\s*(of)?\s*(\d+(?:\.\d+)?)\s*(mm|cm|m|ft)", text):
        add_dim(m.group(1), m.group(3), m.group(4))

    # Volume
    for m in re.finditer(r"(\d+(?:\.\d+)?)\s*(cum|cu\.?m|m3)", text):
        add_dim("volume", m.group(1), m.group(2))

    # Quantity/Count
    for m in re.finditer(r"(\d+)\s*(nos|no|numbers|pcs|pieces|units|sets|points)", text):
        add_dim("quantity", m.group(1), m.group(2))

    return dims

def calculate_overlap_score(input_text, db_text):
    """Calculate token overlap score between input and database text"""
    if not db_text:
        return 0
    
    in_tokens = set(input_text.lower().split())
    db_tokens = set(db_text.lower().split())
    
    # Filter stop words
    stop_words = {'and', 'or', 'with', 'for', 'of', 'in', 'the', 'a', 'to'}
    db_tokens = db_tokens - stop_words
    
    if not db_tokens:
        return 0

    common = in_tokens.intersection(db_tokens)
    return len(common) / len(db_tokens)

def process_raw_text(full_text):
    """
    Main text processing engine using token overlap matching
    Returns list of BOQ items
    """
    results = []
    lines = full_text.split('\n')
    
    current_room = "General"
    last_detected_work = None

    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        norm_line = normalize_text(line)
        
        # 1. CHECK ROOM
        is_room_header = False
        for _, r in room_df.iterrows():
            r_keys = [str(r.get('keyword', '')), str(r.get('room_name', ''))]
            if any(k.lower() in norm_line for k in r_keys if len(k) > 2):
                current_room = r['room_name']
                last_detected_work = None
                is_room_header = True
                break
        
        if is_room_header:
            continue

        # 2. CHECK WORK (TOKEN OVERLAP)
        best_match = None
        best_score = 0.0
        
        for _, w in work_df.iterrows():
            score_kw = calculate_overlap_score(norm_line, str(w.get('keyword', '')))
            score_name = calculate_overlap_score(norm_line, str(w.get('work_name', '')))
            
            max_score = max(score_kw, score_name)
            
            if max_score > best_score and max_score >= 0.5:
                best_score = max_score
                best_match = w
        
        if best_match is not None:
            if last_detected_work:
                results.append(last_detected_work)
            
            last_detected_work = {
                "room_name": current_room,
                "work_code": best_match['work_code'],
                "work_name": best_match['work_name'],
                "dimensions": extract_metrics(line),
                "source": "TEXT_PARSER",
                "confidence": round(best_score, 2)
            }
        else:
            # 3. CHECK FOR DIMENSIONS TO ATTACH
            line_dims = extract_metrics(line)
            if line_dims and last_detected_work:
                for key, val_list in line_dims.items():
                    if key in last_detected_work["dimensions"]:
                        last_detected_work["dimensions"][key].extend(val_list)
                    else:
                        last_detected_work["dimensions"][key] = val_list

    if last_detected_work:
        results.append(last_detected_work)

    return results

# =========================================================
# 4. VISION ENGINE (AI-Powered Image Processing)
# =========================================================

class VisionProcessor:
    def __init__(self):
        self.model = genai.GenerativeModel(GEMINI_MODEL_NAME)

    def scan_image(self, image_path):
        """Process image using Gemini Vision API"""
        if not os.path.exists(image_path):
            return []
        
        try:
            file_ref = genai.upload_file(image_path)
            
            prompt = """
            Act as a Quantity Surveyor. Extract BOQ items from this image.
            Output purely as a JSON list. 
            Format:
            [
              {"room": "Bedroom 1", "desc": "Vitrified Flooring", "qty": "150 sqft"},
              {"room": "Kitchen", "desc": "Wall Dado", "qty": "200 sqft"}
            ]
            Do not use markdown formatting.
            """
            
            res = self.model.generate_content([file_ref, prompt])
            raw_json = res.text.replace("```json", "").replace("```", "").strip()
            return json.loads(raw_json)
        except Exception as e:
            sys.stderr.write(f"Vision processing error: {e}\n")
            return []

    def normalize_to_db(self, ai_data):
        """Map AI-extracted items to database works"""
        final_rows = []
        for item in ai_data:
            desc = item.get("desc", "")
            closest = get_close_matches(desc, WORK_NAMES, n=1, cutoff=0.3)
            
            if closest:
                db_name = closest[0]
                db_code = WORK_MAP.get(db_name, "MISC")
            else:
                db_name = desc
                db_code = "MANUAL_REVIEW"

            final_rows.append({
                "room_name": item.get("room", "General"),
                "work_code": db_code,
                "work_name": db_name,
                "dimensions": extract_metrics(item.get("qty", "")),
                "source": "VISION_ENGINE",
                "confidence": 0.85 if closest else 0.60
            })
        return final_rows

# =========================================================
# 5. FILE EXTRACTION UTILITIES
# =========================================================

def extract_text_from_file(path):
    """Extract text from various file formats"""
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")
    
    ext = os.path.splitext(path)[1].lower()
    
    if ext == ".txt":
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    
    elif ext == ".docx":
        doc = Document(path)
        return "\n".join([p.text for p in doc.paragraphs])
    
    elif ext == ".pdf":
        with pdfplumber.open(path) as pdf:
            return "\n".join([p.extract_text() or "" for p in pdf.pages])
    
    else:
        raise ValueError(f"Unsupported file format: {ext}")

# =========================================================
# 6. MAIN API ENTRY POINT
# =========================================================

def main():
    """
    Main entry point for Node.js integration
    Reads JSON from stdin, processes, outputs JSON to stdout
    """
    try:
        # Read input from stdin
        input_line = sys.stdin.read().strip()
        
        if not input_line:
            output_error("No input received")
            return
        
        input_data = json.loads(input_line)
        mode = str(input_data.get('mode', ''))
        user_input = input_data.get('input', '')
        
        if not mode or not user_input:
            output_error("Missing 'mode' or 'input' in request")
            return
        
        final_data = []
        
        # MODE 1: Raw Text Processing
        if mode == '1':
            final_data = process_raw_text(user_input)
        
        # MODE 2: Document Processing
        elif mode == '2':
            raw_text = extract_text_from_file(user_input)
            final_data = process_raw_text(raw_text)
        
        # MODE 3: Image Processing (Vision Engine)
        elif mode == '3':
            vp = VisionProcessor()
            raw_ai_data = vp.scan_image(user_input)
            final_data = vp.normalize_to_db(raw_ai_data)
        
        else:
            output_error(f"Invalid mode: {mode}")
            return
        
        # Output success response
        output_success(final_data)
        
    except json.JSONDecodeError as e:
        output_error(f"Invalid JSON input: {e}")
    except FileNotFoundError as e:
        output_error(f"File error: {e}")
    except Exception as e:
        output_error(f"Processing error: {e}")

def output_success(data):
    """Output successful result as JSON"""
    result = {
        "success": True,
        "data": data,
        "count": len(data)
    }
    print(json.dumps(result, indent=2))

def output_error(message):
    """Output error as JSON"""
    result = {
        "success": False,
        "error": message,
        "data": [],
        "count": 0
    }
    print(json.dumps(result, indent=2))

# =========================================================
# 7. SCRIPT EXECUTION
# =========================================================

if __name__ == "__main__":
    main()